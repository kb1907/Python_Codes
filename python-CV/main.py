from docx import Document
from docx.shared import Inches

document = Document()

# Profile Picture
document.add_picture("Kaan_Boke photo.png", width=Inches(2.0))

# Personal Info
name = input("What is your name?: ")
email = input("What is your email address?: ")
country = input("Where are you live?: ")
phone = input("What is your phone number: ")
document.add_paragraph(
    name + ' | ' + email + ' | ' + country + ' | ' + phone
)

# Detailed info about me
document.add_heading("About Me")
document.add_paragraph(
    input("Tell me about yourself?:  ")
)

# Work experience
document.add_heading("Work Experience")
while True:
    has_experience = input("Do you have work experience? Yes or No :  ").lower()
    if has_experience == "yes":
        p = document.add_paragraph()
        company = input("Enter company/organization name:  ")
        from_date = input("from date:   ")
        end_date = input("end date:  ")
        p.add_run(company + '\t'+'\t').bold = True
        p.add_run(from_date + ' ' + end_date + '\n').italic = True
        experience_details = input("Tell me about your work experience at  " + company + ' ' + ':')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading("Skills")
while True:
    has_skills = input("Do you want to add the skill? Yes or No :  ").lower()
    if has_skills == "yes":
        p = document.add_paragraph()
        p.style = 'List Bullet'
        skill = input("Enter your skill:  ")
        level_of_skill = input("Level of your skill? Beginner, Intermediate, Professional :   ").lower()
        p.add_run(skill + '\t').bold = True
        p.add_run(level_of_skill + '\n').italic = True
    else:
        break
# Footer
section = document.sections[0]
footer= section.footer
p= footer.paragraphs[0]
p.text = "CV generated using Python"

document.save("Kaan_BOKE_CV.docx")
